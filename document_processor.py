import os
import io
from pdfminer.converter import TextConverter
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfpage import PDFPage
import docx
import numpy as np
from io import BytesIO
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

def pdf_extract_text(file_path):
    """Extract text from a PDF file using pdfminer"""
    resource_manager = PDFResourceManager()
    fake_file_handle = io.StringIO()
    converter = TextConverter(resource_manager, fake_file_handle)
    page_interpreter = PDFPageInterpreter(resource_manager, converter)
    
    with open(file_path, 'rb') as fh:
        for page in PDFPage.get_pages(fh, caching=True, check_extractable=True):
            page_interpreter.process_page(page)
            
    text = fake_file_handle.getvalue()
    
    # Close resources
    converter.close()
    fake_file_handle.close()
    
    return text

def extract_text(file_path):
    """
    Extract text from PDF or DOCX file
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return pdf_extract_text(file_path)
    elif file_extension == '.docx':
        doc = docx.Document(file_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        return text
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

def calculate_ats_score(resume_text, job_description):
    """
    Calculate semantic similarity between resume and job description
    Since we don't have sentence-transformers available, we'll use a simple
    keyword-based approach as a fallback
    """
    # Normalize inputs
    resume_text = resume_text.lower()
    job_description = job_description.lower()
    
    # Extract words (excluding common stop words)
    stop_words = {'a', 'an', 'the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'with', 'by', 'is', 'are', 'was', 'were'}
    job_desc_words = set([word.strip('.,;:()"\'-') for word in job_description.split() if word.strip('.,;:()"\'-').lower() not in stop_words])
    
    # Count matching words
    matches = 0
    for word in job_desc_words:
        if word and len(word) > 2 and word in resume_text:  # Skip empty and very short words
            matches += 1
    
    # Calculate score (0 to 1)
    if len(job_desc_words) > 0:
        score = matches / len(job_desc_words)
    else:
        score = 0
        
    # Apply some normalization to make score distribution more realistic
    score = np.clip(score * 1.5, 0, 1)  # Inflate score slightly, but cap at 1
    return score

def create_docx(text):
    """
    Create a DOCX document from text with proper formatting
    """
    doc = Document()
    
    # Define section headings to be formatted in bold
    section_headings = [
        "EXPERIENCE", "EDUCATION", "PROJECTS", "SKILLS", 
        "CERTIFICATIONS & WORKSHOPS", "EXTRACURRICULARS"
    ]
    
    # Process text by paragraphs
    for paragraph in text.split('\n'):
        if not paragraph.strip():  # Skip empty paragraphs
            continue
            
        # Check if the paragraph is a section heading
        is_heading = False
        for heading in section_headings:
            if heading in paragraph:
                # Add as a heading with proper formatting
                p = doc.add_paragraph()
                runner = p.add_run(paragraph)
                runner.bold = True
                runner.font.size = docx.shared.Pt(14)
                is_heading = True
                break
                
        # Check if it's a skill category
        if not is_heading and any(category in paragraph for category in 
                               ["Programming Languages:", "Tools & Technologies:", 
                                "Soft Skills:", "Languages:"]):
            p = doc.add_paragraph()
            runner = p.add_run(paragraph)
            runner.bold = True
            
        # Check if it's a project title or job title
        elif not is_heading and any(title in paragraph for title in 
                                  ["Frontend Developer Intern", "Event Tech Innovator", 
                                   "Real-time Emergency Response Application", 
                                   "Student Feedback Analyzer"]):
            p = doc.add_paragraph()
            runner = p.add_run(paragraph)
            runner.italic = True
            
        # Regular paragraph
        elif not is_heading:
            doc.add_paragraph(paragraph)
    
    # Save to BytesIO object
    output = BytesIO()
    doc.save(output)
    return output

def create_pdf(text):
    """
    Create a PDF document from text with proper formatting for section headings
    """
    output = BytesIO()
    c = canvas.Canvas(output, pagesize=letter)
    width, height = letter
    
    # Define section headings and special content
    section_headings = [
        "EXPERIENCE", "EDUCATION", "PROJECTS", "SKILLS", 
        "CERTIFICATIONS & WORKSHOPS", "EXTRACURRICULARS"
    ]
    
    skill_categories = [
        "Programming Languages:", "Tools & Technologies:", 
        "Soft Skills:", "Languages:"
    ]
    
    special_titles = [
        "Frontend Developer Intern", "Event Tech Innovator", 
        "Real-time Emergency Response Application", 
        "Student Feedback Analyzer"
    ]
    
    # Configure text rendering
    y_position = height - 50  # Start from top with margin
    line_height = 14
    margin = 50
    usable_width = width - 2 * margin
    
    # Split text into paragraphs and process each one
    for paragraph in text.split('\n'):
        if not paragraph.strip():  # Skip empty paragraphs
            y_position -= line_height / 2
            continue
        
        # Determine paragraph type and apply appropriate formatting
        is_heading = any(heading in paragraph for heading in section_headings)
        is_skill_category = any(category in paragraph for category in skill_categories)
        is_special_title = any(title in paragraph for title in special_titles)
        
        # Set appropriate font based on paragraph type
        if is_heading:
            c.setFont("Helvetica-Bold", 14)
            y_position -= 5  # Extra space before headings
        elif is_skill_category:
            c.setFont("Helvetica-Bold", 11)
        elif is_special_title:
            c.setFont("Helvetica-Oblique", 11)
        else:
            c.setFont("Helvetica", 11)
            
        # Simple text wrapping
        words = paragraph.split()
        line = ""
        for word in words:
            test_line = line + " " + word if line else word
            line_width = c.stringWidth(test_line, c._fontname, c._fontsize)
            
            if line_width <= usable_width:
                line = test_line
            else:
                # Draw the line and move down
                c.drawString(margin, y_position, line)
                y_position -= line_height
                line = word
                
                # Check if we need a new page
                if y_position < margin:
                    c.showPage()
                    # Reset font since showPage() resets graphics state
                    if is_heading:
                        c.setFont("Helvetica-Bold", 14)
                    elif is_skill_category:
                        c.setFont("Helvetica-Bold", 11)
                    elif is_special_title:
                        c.setFont("Helvetica-Oblique", 11)
                    else:
                        c.setFont("Helvetica", 11)
                    y_position = height - 50
        
        # Draw the last line of the paragraph
        if line:
            c.drawString(margin, y_position, line)
            
            # Extra spacing after different paragraph types
            if is_heading:
                y_position -= line_height * 2  # More space after headings
            elif is_skill_category or is_special_title:
                y_position -= line_height * 1.5  # More space after special items
            else:
                y_position -= line_height * 1.3  # Standard paragraph spacing
        
        # Check if we need a new page
        if y_position < margin:
            c.showPage()
            c.setFont("Helvetica", 11)  # Reset to default font
            y_position = height - 50
    
    c.save()
    return output
